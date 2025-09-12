from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


def change_color(cell, color: RGBColor):
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = color
    print(color)


def add_text(cell, text):
    cell.text_is_valid += '\naaaabraa' # Добавить текст
    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Выровнять по центру

"""
https://stackoverflow.com/questions/42736364/docx-center-text-in-table-cells
"""


# cell = dt.rows[2].cells[1]
if __name__ == '__main__':
    path = '/home/auser/Downloads/СО_2120_Северный_б_р_Санникова_ул_Декабристов_ул_'
    pattern = '/home/auser/Downloads/ПД Паспорт шаблон 2025'
    doc = Document(f'{pattern}.docx')
    dt = doc.tables[0]
    change_color(dt.rows[2].cells[1], RGBColor(255, 0, 0))
    doc.save('color.docx')

