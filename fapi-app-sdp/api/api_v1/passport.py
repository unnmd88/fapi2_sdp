from collections.abc import Sequence
from fastapi.concurrency import run_in_threadpool

from fastapi import APIRouter, HTTPException
from fastapi import File, UploadFile
from starlette import status
from fastapi.responses import FileResponse

from docx import Document
from core.config import UPLOADS_URL
from sdp_lib.passport.api import create_passport
from utils.files import PassportSaver

router = APIRouter(tags=['Passport'])
DIR_NAME = UPLOADS_URL / 'passport'


@router.post("/validation")
async def upload_docx(
    # err_color: str = Form(...),
    files: Sequence[UploadFile] = File(...),
) :
    # passports = [create_passport(src_file.filename, Document(src_file.file)) for src_file in files]

    if not files:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='Не предоставлено ни одного файла для обработки.'
        )
    allowed_files, bad_files = [], []
    for f in files:
        try:
            doc = Document(f.file)
            allowed_files.append(create_passport(f.filename, Document(f.file)))
        except Exception:
            bad_files.append(f)
    if bad_files:
        raise HTTPException(
            400,
            detail=f'Invalid document type. Files: {", ".join(f.filename for f in bad_files)}'
        )
    # for f in files:
    #     display_all_tables(Document(f.file))
    passport_saver = PassportSaver(allowed_files)
    path = await run_in_threadpool(passport_saver.save_and_get_as_docx_if_only_one_else_as_zip_archive,)
    filename = 'result.docx' if path.suffix == ".docx" else path.name
    return FileResponse(path, headers={'Content-Disposition': f'attachment; filename={filename}'})

    # bstream.seek(0)
    # return StreamingResponse(
    #     bstream,
    #     headers=headers,
    #     media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # )



# @router.post("/passport-validation")
# async def upload_docx(
#     err_color: str = Form(...),
#     files: Sequence[UploadFile] = File(...),
# ) :
#     # passports = [create_passport(Document(src_file.file)) for src_file in files]
#     for src_file in files:
#         doc = Document(src_file.file)
#         print(type(src_file.file))
#         bstream = io.BytesIO()
#         passp = create_passport(doc)
#         passp.get_docx().save(bstream)
#     headers = {
#         'Content-Disposition': 'attachment; filename="filename1.docx"'
#     }
#     bstream.seek(0)
#     return StreamingResponse(
#         bstream,
#         headers=headers,
#         media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#     )
#
#     # for src_file in files:
#     #     doc = docx.Document(src_file.file)
#     #     passp = create_passport(doc)
#     #     passp.get_docx().save('cccaddaaa.docx')
#     # headers = {'Content-Disposition': 'attachment; filename="zagruzi_a.docx"'}
#     # return FileResponse(
#     #     'cccaddaaa.docx', headers=headers
#     # )
#     # for src_file in files:
#     #     doc = docx.Document(src_file.file)
#     #     fs = io.BytesIO()
#     #     doc.save(fs)
#     #     passports.append(fs)
#     #     fs.seek(0)

