import itertools
import re
from collections.abc import Generator, Sequence, Iterable
from enum import IntEnum
from typing import Any

from docx import Document
from docx.shared import  RGBColor
from docx.table import (
    _Rows,
    Table
)

from sdp_lib.passport.constants import (
    allowed_column_lengths_dt,
    allowed_min_num_rows,
    DirectionEntities,
    PatternsDirectionTable,
    AllowedValues,
    mapping_direction_data,
    DirectionDataContainer,
    head_rows_data_dt, Patterns
)
from sdp_lib.passport.base import (
    DirectionDataRow,
    CellData,
    CellMapping,
    MessageStorage, TheTable
)
from sdp_lib.passport.utils import (
    found_pos_start_num,
    repair_string_if_sep_in_illegal_pos
)
from sdp_lib.passport.validation.common_validators import (
    validate_geometry,
    validate_sequence_directions_or_stages_nums_and_create_cell,
    match_one_string_to_many_patterns_and_get_alias_and_create_cell,
    create_cells_for_head_row,
    num_validate_and_create_cell,
    lrstrip_in_cell_and_create_cell_mappings,
    gen_default_cells,
    create_default_cell
)
from sdp_lib.passport.text_messages import Text
from sdp_lib.utils_common.utils_common import timed


entity_patterns_and_aliases = (
    (PatternsDirectionTable.vehicle.value, DirectionEntities.vehicle, ),
    (PatternsDirectionTable.arrow.value, DirectionEntities.arrow, ),
    (PatternsDirectionTable.pedestrian.value, DirectionEntities.pedestrian, ),
    (PatternsDirectionTable.always_red.value, DirectionEntities.always_red, ),
    (PatternsDirectionTable.public.value, DirectionEntities.public,),
    (PatternsDirectionTable.tram.value, DirectionEntities.tram, ),
    (PatternsDirectionTable.velo.value, DirectionEntities.velo, ),
)


def validate_traffic_lights(
    cell_mapping: CellMapping,
    direction_entity: DirectionEntities,
    tl_patterns: Sequence[str | re.Pattern],
) -> CellData:
    ms = MessageStorage([], [])
    src_txt = cell_mapping.cell.text # Пример: "Тр. 7,8,9,10"
    repaired_text = repair_string_if_sep_in_illegal_pos(src_txt, ',')
    tlc_entity_is_valid = tlc_nums_is_valid = True
    if len(repaired_text) != len(src_txt):
        ms.add_errors(Text.illegal_pos_for_char(','))
    if (pos_start_nums_tlc:= found_pos_start_num(repaired_text)) is None:
        tlc_entity, nums = repaired_text, ''
        ms.add_errors(Text.has_not_num_tlc)
        tlc_nums_is_valid = False
    else:
        tlc_entity = repaired_text[:pos_start_nums_tlc] # Из примера: "Тр. "
        nums = repaired_text[pos_start_nums_tlc:].replace(' ', '') # Из примера: "7,8,9,10"
    split_nums = nums.split(',') if nums else ''
    if bad_nums := [n for n in split_nums if not re.sub(Patterns.s_char.value, '', n).isdigit()]:
        tlc_nums_is_valid = False
        ms.add_errors(Text.invalid_nums(bad_nums))
    if direction_entity is not None:
        if not any(re.search(p, tlc_entity) for p in tl_patterns):
            tlc_entity_is_valid = False
            ms.add_errors(Text.invalid_type_tlc)
    else:
        tlc_entity_is_valid = False
    res = bool(tlc_entity_is_valid and tlc_nums_is_valid)
    return CellData(
        value=src_txt,
        text_is_valid=res,
        context_is_valid=res,
        recovered_val=repaired_text,
        messages=ms,
        cell_mapping=cell_mapping
    )


def validate_timings(
    cell_mapping: CellMapping,
    direction_entity: DirectionEntities,
    timings: AllowedValues,
) -> CellData:
    ms = MessageStorage([], [])
    txt = cell_mapping.cell.text
    text_is_valid = False
    try:
        val_i = int(txt)
        text_is_valid = True
        if direction_entity is None:
            return CellData(
                txt, text_is_valid, None, converted_val=val_i, messages=ms, cell_mapping=cell_mapping
            )
    except ValueError:
        ms.add_errors(Text.is_not_a_number)
        return CellData(txt, text_is_valid, text_is_valid, messages=ms, cell_mapping=cell_mapping)
    if timings.min <= val_i <= timings.max:  # OK case
        return CellData(txt, text_is_valid, text_is_valid, converted_val=val_i, messages=ms, cell_mapping=cell_mapping)
    if val_i < timings.min:
        err = Text.val_must_be_gt(timings.min)
    elif val_i > timings.max:
        err = Text.val_must_be_lt(timings.max)
    else:
        raise Exception(f'Debug: val_to_validate not fully validated')
    ms.add_errors(err)
    return CellData(txt, text_is_valid, False, messages=ms, cell_mapping=cell_mapping)


def validate_always_red_col_and_create_cell_data(
    cell_mapping: CellMapping,
    direction_entity: DirectionEntities,
):
    if direction_entity is None:
        return create_default_cell(cell_mapping)
    ms = MessageStorage([], [])
    src_txt = cell_mapping.cell.text
    if ((direction_entity != DirectionEntities.always_red and src_txt == '')
        or
        (direction_entity == DirectionEntities.always_red and re.match(PatternsDirectionTable.always_red_text_yes.value, src_txt))
    ): # ОК
        return CellData(
            value=src_txt,
            text_is_valid=True,
            context_is_valid=True,
            messages=ms,
            cell_mapping=cell_mapping
        )
    red_pattern = re.match(PatternsDirectionTable.always_red_text_yes.value, src_txt)
    if ((direction_entity != DirectionEntities.always_red and red_pattern)
        or
        (src_txt and red_pattern is None)
    ):
        ms.add_errors(Text.invalid_value)
    return CellData(
        value=src_txt,
        text_is_valid=bool(src_txt == '' or red_pattern is not None),
        context_is_valid=bool(ms.errors),
        messages=ms,
        cell_mapping=cell_mapping
    )


def validate_toov(
    cell_mapping: CellMapping,
    direction_entity: DirectionDataContainer,
    toov_patterns: Sequence[str | re.Pattern]
) -> CellData:
    if direction_entity is None:
        return create_default_cell(cell_mapping)
    txt = cell_mapping.cell.text
    text_is_valid = any(re.match(p, txt) is not None for p in toov_patterns)
    return CellData(
        value=txt,
        text_is_valid=text_is_valid,
        context_is_valid=False if not text_is_valid else None,
        cell_mapping=cell_mapping,
        messages=MessageStorage([Text.invalid_value_for_direction_entity] if not text_is_valid else [], [])
    )


class RowPosition(IntEnum):
    num                 = 0
    entity              = 1
    stages              = 2
    tlc                 = 3
    t_green_ext         = 4
    t_green_flashing    = 5
    t_green_yellow      = 6
    t_red               = 7
    t_red_yellow        = 8
    t_z                 = 9
    t_zz                = 10
    always_red          = 11
    toov_red            = 12
    toov_green          = 13
    description         = 14


@timed
def validate_and_create_directions_table(i_table: int, table: Table, ) -> TheTable:
    table_rows: _Rows = table.rows
    the_table = TheTable(
        i_table,
        table,
        validate_geometry(table, allowed_column_lengths_dt, allowed_min_num_rows)
    )
    # Первый этап валидации - проверка корректности геометрии таблицы:
    # Количество столбцов и минимальное количество строк.
    if not the_table.geometry_check_list.is_valid:
        for err in the_table.geometry_check_list.get_errors():
            the_table.messages.add_errors(err)
        target = table.add_row()
        target.cells[0].text = '\n'.join(f"*{e}" for e in the_table.messages.chain())
        target.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        return table
    first_row = lrstrip_in_cell_and_create_cell_mappings(i_table, 0, table_rows[0].cells)
    first_row = tuple(gen_default_cells(first_row[0]))
    second_row = tuple(create_cells_for_head_row(
        i_table, 1, table_rows[1].cells, head_rows_data_dt.row1_col_patterns, head_rows_data_dt.row1_col_names
    ))
    the_table.load_head_rows((DirectionDataRow(first_row), DirectionDataRow(second_row)))
    # Второй этап валидации - проверка корректности названий колонок:
    if not the_table.head_rows[1].is_valid:
        return table
    rows, empty_rows = [], []
    for i in range(2, len(table_rows)):
        cell_mappings, empty_cells = lrstrip_in_cell_and_create_cell_mappings(i_table, i, table_rows[i].cells)
        if len(cell_mappings) == empty_cells:
            empty_rows.append(DirectionDataRow(tuple(gen_default_cells(c) for c in cell_mappings)))
        else:
            num = num_validate_and_create_cell(cell_mappings[0]).write_messages_to_table_cell()
            entity = match_one_string_to_many_patterns_and_get_alias_and_create_cell(
                cell_mappings[1],
                entity_patterns_and_aliases,
                True,
            ).write_messages_to_table_cell()
            direction_data: DirectionDataContainer = mapping_direction_data.get(entity.converted_val)
            entity_name = direction_data.entity if direction_data is not None else None
            stages = validate_sequence_directions_or_stages_nums_and_create_cell(
                cell_mappings[2]
            ).write_messages_to_table_cell()
            tl_patterns = direction_data.tl_patterns if direction_data is not None else None
            tl = validate_traffic_lights(cell_mappings[3], entity_name, tl_patterns).write_messages_to_table_cell()
            timing_cells_iterator: Iterable = direction_data.get_timings() if direction_data is not None else range(7)
            timings: Generator[CellData, Any, None] = (
                validate_timings(
                    cell_mappings[ii],
                    entity_name,
                    col_name,
                ).write_messages_to_table_cell()
                for ii, col_name in enumerate(timing_cells_iterator, 4)
            )
            always_red = validate_always_red_col_and_create_cell_data(
                cell_mappings[11], entity_name
            ).write_messages_to_table_cell()
            toov_patterns = direction_data.toov_patterns if direction_data is not None else None
            toov_red = validate_toov(
                cell_mappings[12], direction_data, toov_patterns
            ).write_messages_to_table_cell()
            toov_green = validate_toov(
                cell_mappings[13], direction_data, toov_patterns
            ).write_messages_to_table_cell()
            description = create_default_cell(cell_mappings[14])
            chain = itertools.chain(
                (num, entity, stages, tl),
                timings,
                (always_red, toov_red, toov_green, description),
            )
            rows.append(DirectionDataRow(tuple(c for c in chain)))
    the_table.load_data_rows(rows)
    the_table.load_empty_rows(empty_rows)
    return the_table

if __name__ == '__main__':

    path4 = "C:\Programms\py.projects\sdp_lib\sdp_lib\passport\СО_2094_ул_Островитянова_ул_Ак_Волгина (2).docx"
    path5 = '/home/auser/py.projects/sdp_lib/sdp_lib/passport/СО_2094_ул_Островитянова_ул_Ак_Волгина_2.docx'
    path6 = '/home/auser/py.projects/sdp_lib/sdp_lib/passport/passport2/validation/ПД Паспорт шаблон 2025.docx'

    doc = Document(path5)
    # c = CheckListTable()
    validate_and_create_directions_table(0, doc.tables[0])

