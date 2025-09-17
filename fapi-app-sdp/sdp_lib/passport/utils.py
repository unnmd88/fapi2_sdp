import re
from collections.abc import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.table import _Cell

from sdp_lib.passport.constants import Patterns
from sdp_lib.passport.base import CellData


RGB_RED = RGBColor(255, 0, 0)


def remove_left_light_spaces_from_cell_text(cell: _Cell) ->_Cell:
    cell.text = cell.text.lstrip().rstrip()
    return cell


def repair_string_if_sep_in_illegal_pos(
    string: str,
    sep=',',
) -> str:
    if string:
        string = re.sub(Patterns.several_commas.value if sep == ',' else sep, sep, string)
        return re.sub(Patterns.comma_is_start_end_or_end.value if sep == ',' else sep, '', string)
    return  string


def add_text_co_cell(cell: _Cell, messages: Iterable[str], text_color: RGBColor | tuple[int, int, int] = None):
    new_line = '\n'
    cell.text += f'{new_line}{new_line.join(f"*{t}" for t in messages)}'
    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text_color:
        cell.paragraphs[0].runs[0].font.color.rgb = text_color if isinstance(text_color, RGBColor) else RGBColor(text_color)


def write_messages_to_cell(*cells: CellData, color: RGBColor = None, sep='\n'):
    for cell in cells:
        new_txt = sep.join(f'*{m}' for m in cell.messages.chain())
        if new_txt:
            cell.cell_mapping.cell.text = f'{cell.cell_mapping.cell.text}{sep}{new_txt}'
            color = color or RGB_RED
            para = cell.cell_mapping.cell.paragraphs[0]
            para.runs[0].font.color.rgb = color
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def found_pos_start_num(src: Iterable[str]) -> int | None:
    for i, char in enumerate(src):
        if char.isdigit():
            return i
    return None


def display_all_tables(doc_x):
    """ Выводит на экран данные всех таблиц doc(x) файла. """
    for i, table in enumerate(doc_x.tables, 1):
        print(f'-- Start Table {i} --')
        # print(f'Столбцов: {len(table.rows[0].cells)} | Строк: {len(table.rows)}')
        print(f'Столбцов: {len(table.columns)} | Строк: {len(table.rows)}')
        for ii, row in enumerate(table.rows):
            print(f'{ii}: {[cell.text for cell in row.cells]}')
        print(f'-- End Table {i} --')
        print(f'*' * 100)


if __name__ == '__main__':
    path = '/home/auser/Downloads/СО_2120_Северный_б_р_Санникова_ул_Декабристов_ул_'
    pattern = '/home/auser/Downloads/ПД Паспорт шаблон 2025'
    path_sdp = "C:\Programms\py.projects\sdp_lib\sdp_lib\passport\СО_2094_ул_Островитянова_ул_Ак_Волгина (2)"
    path5 = '/home/auser/py.projects/sdp_lib/sdp_lib/passport/СО_2094_ул_Островитянова_ул_Ак_Волгина_2.docx'
    path7 = '/home/auser/py.projects/sdp_lib/sdp_lib/passport/СО_72_Тургеневская_пл_Мясницкая_ул_Сретенский_б_р_Чистопрудный_б.docx'

    doc = Document(path7)
    display_all_tables(doc)
